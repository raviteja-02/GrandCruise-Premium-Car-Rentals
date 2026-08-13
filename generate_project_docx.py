import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_code_block(doc, code):
    """Add a code block to the document"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    paragraph.style = 'No Spacing'

def update_section_8(doc):
    """Update Section 8: Dataset Description with comprehensive Django models information"""
    
    # Find the section 8 heading
    section_8_found = False
    for i, paragraph in enumerate(doc.paragraphs):
        if "8. DATASET DESCRIPTION" in paragraph.text or "8. Dataset Description" in paragraph.text:
            section_8_found = True
            # Remove all paragraphs from this point until the next section
            paragraphs_to_remove = []
            for j in range(i + 1, len(doc.paragraphs)):
                next_para = doc.paragraphs[j]
                if next_para.text.strip().startswith('9.') or next_para.text.strip().startswith('## 9'):
                    break
                paragraphs_to_remove.append(next_para)
            
            # Remove the paragraphs
            for para in paragraphs_to_remove:
                p = para._element
                p.getparent().remove(p)
            
            # Add the new content after the heading
            new_content = [
                "The dataset for this car booking system consists of comprehensive information about vehicles, users, bookings, and system configurations. The data is structured across multiple Django models that form the backbone of the application's database schema.",
                "",
                "8.1 Data Sources",
                "",
                "The dataset is primarily generated from:",
                "• User registrations and profile information",
                "• Car inventory management",
                "• Booking transactions", 
                "• System configuration settings",
                "",
                "8.2 Data Structure",
                "",
                "The dataset is organized into the following main categories with detailed model specifications:",
                "",
                "8.2.1 User Data Models",
                "",
                "User (Django Built-in):",
                "• username: CharField - Unique username for login (max_length=150, unique=True)",
                "• email: EmailField - User's email address (unique=True)",
                "• first_name: CharField - User's first name (max_length=150)",
                "• last_name: CharField - User's last name (max_length=150)",
                "• is_active: BooleanField - Account activation status (default=True)",
                "• date_joined: DateTimeField - Account creation timestamp (auto_now_add=True)",
                "",
                "UserProfile:",
                "• user: OneToOneField - Link to User model (CASCADE delete)",
                "• phone_number: CharField - Contact phone number (max_length=15, blank=True)",
                "• address: TextField - User's address (blank=True)",
                "• created_at: DateTimeField - Profile creation time (auto_now_add=True)",
                "• updated_at: DateTimeField - Last update timestamp (auto_now=True)",
                "",
                "8.2.2 Vehicle Data Models",
                "",
                "Car:",
                "• brand: CharField - Car manufacturer (max_length=100)",
                "• model: CharField - Car model name (max_length=100)",
                "• year: IntegerField - Manufacturing year",
                "• seats: IntegerField - Number of seats",
                "• price_per_day: DecimalField - Daily rental rate (max_digits=10, decimal_places=2)",
                "• image: URLField - Main car image URL (max_length=500)",
                "• description: TextField - Car description",
                "• fuelType: CharField - Fuel type (max_length=50, blank=True)",
                "• transmission: CharField - Transmission type (max_length=50, blank=True)",
                "• features: JSONField - Car features list (default=list, blank=True)",
                "• specifications: JSONField - Technical specifications (default=list, blank=True)",
                "• mileage: IntegerField - Vehicle mileage (default=0, blank=True)",
                "• engine: CharField - Engine details (max_length=100, blank=True)",
                "• color: CharField - Car color (max_length=100, blank=True)",
                "• category: CharField - Car category (max_length=100, blank=True)",
                "• rating: FloatField - Average rating (default=0.0)",
                "• reviews: IntegerField - Number of reviews (default=0)",
                "• is_available: BooleanField - Availability status (default=True)",
                "• created_at: DateTimeField - Creation timestamp (auto_now_add=True)",
                "• updated_at: DateTimeField - Update timestamp (auto_now=True)",
                "",
                "CarGallery:",
                "• car: OneToOneField - Link to Car model (CASCADE delete)",
                "• gallery_url: URLField - Gallery URL (max_length=500, default='')",
                "• images: JSONField - Image URLs list (default=list, blank=True)",
                "",
                "8.2.3 Booking Data Models",
                "",
                "Booking:",
                "• user: ForeignKey - Link to User model (CASCADE delete)",
                "• car: ForeignKey - Link to Car model (CASCADE delete)",
                "• start_date: DateField - Booking start date",
                "• end_date: DateField - Booking end date",
                "• total_price: DecimalField - Total booking cost (max_digits=10, decimal_places=2)",
                "• status: CharField - Booking status (choices: pending/confirmed/cancelled/completed)",
                "• created_at: DateTimeField - Booking creation time (auto_now_add=True)",
                "• updated_at: DateTimeField - Last update timestamp (auto_now=True)",
                "",
                "8.2.4 System Configuration Models",
                "",
                "SystemSettings:",
                "• maintenance_mode: BooleanField - System maintenance status (default=False)",
                "• booking_confirmation_required: BooleanField - Require booking confirmation (default=True)",
                "• max_bookings_per_user: IntegerField - Maximum bookings per user (default=3)",
                "• default_booking_duration: IntegerField - Default booking hours (default=24)",
                "• updated_at: DateTimeField - Settings update time (auto_now=True)",
                "",
                "8.3 Data Relationships",
                "",
                "The dataset maintains the following key relationships:",
                "• User ↔ UserProfile: One-to-one relationship for extended user information",
                "• User ↔ Booking: One-to-many relationship for user bookings",
                "• Car ↔ CarGallery: One-to-one relationship for car image galleries",
                "• Car ↔ Booking: One-to-many relationship for car bookings",
                "• SystemSettings: Singleton pattern for global system configuration",
                "",
                "8.4 Data Quality Features",
                "",
                "The dataset maintains high quality through:",
                "• Input Validation: Field-level constraints and validation",
                "• Automatic Timestamps: Created and updated timestamps for all models",
                "• Relationship Integrity: Foreign key constraints with CASCADE delete",
                "• Data Normalization: Proper model relationships and field types",
                "• Status Tracking: Booking status management with business logic",
                "• Availability Management: Automatic car availability updates",
                "• JSON Fields: Flexible storage for features and specifications",
                "",
                "8.5 Data Volume and Growth",
                "",
                "The dataset is designed to handle:",
                "• User Growth: Scalable user registration and profile management",
                "• Inventory Expansion: Dynamic car catalog with detailed specifications",
                "• Booking Volume: High-frequency booking transactions with status tracking",
                "• System Configuration: Centralized settings management for application behavior",
                "",
                "8.6 Data Security and Privacy",
                "",
                "The dataset implements several security measures:",
                "• User Authentication: Secure user login and session management",
                "• Data Encryption: Sensitive information protection",
                "• Access Control: Role-based permissions and authorization",
                "• Audit Trail: Comprehensive logging of data changes and access patterns"
            ]
            
            # Add the new content
            for content in new_content:
                if content.startswith('8.'):
                    # Add subheading
                    p = doc.add_paragraph()
                    p.add_run(content)
                    p.style = 'Heading 3'
                elif content.startswith('•'):
                    # Add bullet point
                    p = doc.add_paragraph()
                    p.add_run(content)
                    p.style = 'List Bullet'
                elif content.strip() == "":
                    # Add empty paragraph
                    doc.add_paragraph()
                else:
                    # Add regular paragraph
                    p = doc.add_paragraph()
                    p.add_run(content)
            
            break
    
    if not section_8_found:
        print("Section 8 not found in document")

def update_section_9(doc):
    """Update Section 9: Testing with original content"""
    
    # Find the section 9 heading
    section_9_found = False
    for i, paragraph in enumerate(doc.paragraphs):
        if "9. TESTING" in paragraph.text or "9. Testing" in paragraph.text:
            section_9_found = True
            # Remove all paragraphs from this point until the next section
            paragraphs_to_remove = []
            for j in range(i + 1, len(doc.paragraphs)):
                next_para = doc.paragraphs[j]
                if next_para.text.strip().startswith('11.') or next_para.text.strip().startswith('## 11'):
                    break
                paragraphs_to_remove.append(j)
            
            # Remove paragraphs in reverse order
            for j in reversed(paragraphs_to_remove):
                p = doc.paragraphs[j]
                p._element.getparent().remove(p._element)
            
            # Add comprehensive testing content
            add_paragraph(doc, "The purpose of testing is to discover errors. Testing is the process of trying to discover every conceivable fault or weakness in a work product. It provides a way to check the functionality of components, sub assemblies, assemblies and/or a finished product. It is the process of exercising software with the intent of ensuring that the software system meets its requirements and user expectations and does not fail in an unacceptable manner. There are various types of test. Each test type addresses a specific testing requirement.")
            
            add_heading(doc, "9.1 TYPES OF TESTS", 2)
            
            add_heading(doc, "Unit testing", 3)
            add_paragraph(doc, "Unit testing involves the design of test cases that validate that the internal program logic is functioning properly, and that program inputs produce valid outputs. All decision branches and internal code flow should be validated. It is the testing of individual software units of the application. It is done after the completion of an individual unit before integration. This is a structural testing, that relies on knowledge of its construction and is invasive. Unit tests perform basic tests at component level and test a specific business process, application, and/or system configuration. Unit tests ensure that each unique path of a business process performs accurately to the documented specifications and contains clearly defined inputs and expected results.")
            
            add_heading(doc, "Integration testing", 3)
            add_paragraph(doc, "Integration tests are designed to test integrated software components to determine if they actually run as one program. Testing is event driven and is more concerned with the basic outcome of screens or fields. Integration tests demonstrate that although the components were individually satisfaction, as shown by successfully unit testing, the combination of components is correct and consistent. Integration testing is specifically aimed at exposing the problems that arise from the combination of components.")
            
            add_heading(doc, "Functional test", 3)
            add_paragraph(doc, "Functional tests provide systematic demonstrations that functions tested are available as specified by the business and technical requirements, system documentation, and user manuals.")
            add_paragraph(doc, "Functional testing is centered on the following items:")
            add_paragraph(doc, "• Valid Input: identified classes of valid input must be accepted.")
            add_paragraph(doc, "• Invalid Input: identified classes of invalid input must be rejected.")
            add_paragraph(doc, "• Functions: identified functions must be exercised.")
            add_paragraph(doc, "• Output: identified classes of application outputs must be exercised.")
            add_paragraph(doc, "• Systems/Procedures: interfacing systems or procedures must be invoked.")
            add_paragraph(doc, "Organization and preparation of functional tests is focused on requirements, key functions, or special test cases. In addition, systematic coverage pertaining to identify Business process flows; data fields, predefined processes, and successive processes must be considered for testing. Before functional testing is complete, additional tests are identified and the effective value of current tests is determined.")
            
            add_heading(doc, "White Box Testing", 3)
            add_paragraph(doc, "White Box Testing is a testing in which the software tester has knowledge of the inner workings, structure and language of the software, or at least its purpose. It is used to test areas that cannot be reached from a black box level.")
            
            add_heading(doc, "Black Box Testing", 3)
            add_paragraph(doc, "Black Box Testing is testing the software without any knowledge of the inner workings, structure or language of the module being tested. Black box tests, as most other kinds of tests, must be written from a definitive source document, such as specification or requirements document. It is a testing in which the software under test is treated, as a black box. You cannot \"see\" into it. The test provides inputs and responds to outputs without considering how the software works.")
            
            add_heading(doc, "Top Down Integration", 3)
            add_paragraph(doc, "This method is an incremental approach to the construction of program structure. Modules are integrated by moving downward through the control hierarchy, beginning with the main program module. The module subordinates to the main program module are incorporated into the structure in either a depth first or breadth first manner.")
            add_paragraph(doc, "In this method, the software is tested from main module and individual stubs are replaced when the test proceeds downwards.")
            
            add_heading(doc, "Bottom-up Integration", 3)
            add_paragraph(doc, "This method begins the construction and testing with the modules at the lowest level in the program structure. Since the modules are integrated from the bottom up, processing required for modules subordinate to a given level is always available and the need for stubs is eliminated. The bottom up integration strategy may be implemented with the following steps:")
            add_paragraph(doc, "• The low-level modules are combined into clusters that perform a specific software sub-function.")
            add_paragraph(doc, "• A driver (i.e.) the control program for testing is written to coordinate test case input and output.")
            add_paragraph(doc, "• The cluster is tested.")
            add_paragraph(doc, "• Drivers are removed and clusters are combined moving upward in the program structure")
            add_paragraph(doc, "The bottom up approaches tests each module individually and then each module is integrated with a main module and tested for functionality.")
            break
    
    if not section_9_found:
        print("Section 9 heading not found")

def update_section_10(doc):
    """Update Section 10: Source Code with complete file code"""
    
    # Find the section 10 heading
    section_10_found = False
    for i, paragraph in enumerate(doc.paragraphs):
        if "10. SOURCE CODE" in paragraph.text or "10. Source Code" in paragraph.text:
            section_10_found = True
            # Remove all paragraphs from this point until the next section
            paragraphs_to_remove = []
            for j in range(i + 1, len(doc.paragraphs)):
                next_para = doc.paragraphs[j]
                if next_para.text.strip().startswith('11.') or next_para.text.strip().startswith('## 11'):
                    break
                paragraphs_to_remove.append(j)
            
            # Remove paragraphs in reverse order
            for j in reversed(paragraphs_to_remove):
                p = doc.paragraphs[j]
                p._element.getparent().remove(p._element)
            
            # Add comprehensive source code content
            add_heading(doc, "11.1 Django Models", 2)
            add_paragraph(doc, "users/models.py")
            add_code_block(doc, """from django.db import models
from django.contrib.auth.models import User

class UserProfile(models.Model):
    user = models.OneToOneField(User, on_delete=models.CASCADE)
    phone_number = models.CharField(max_length=15, blank=True, null=True)
    address = models.TextField(blank=True, null=True)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    def __str__(self):
        return f"{self.user.username}'s profile"

    class Meta:
        verbose_name = "User Profile"
        verbose_name_plural = "User Profiles"
""")

            add_paragraph(doc, "cars/models.py")
            add_code_block(doc, """from django.db import models
from django.core.validators import MinValueValidator, MaxValueValidator

class CarCategory(models.Model):
    name = models.CharField(max_length=100)
    description = models.TextField(blank=True)
    
    def __str__(self):
        return self.name
    
    class Meta:
        verbose_name_plural = "Car categories"

class CarColor(models.Model):
    name = models.CharField(max_length=50)
    hex_code = models.CharField(max_length=7, default="#000000")
    
    def __str__(self):
        return self.name

class CarEngine(models.Model):
    name = models.CharField(max_length=100)
    fuel_type = models.CharField(max_length=50)
    displacement = models.CharField(max_length=50)
    power = models.CharField(max_length=50)
    
    def __str__(self):
        return f"{self.name} - {self.fuel_type}"

class CarFeature(models.Model):
    name = models.CharField(max_length=100)
    description = models.TextField(blank=True)
    
    def __str__(self):
        return self.name

class Car(models.Model):
    brand = models.CharField(max_length=100)
    model = models.CharField(max_length=100)
    year = models.IntegerField()
    seats = models.IntegerField()
    price_per_day = models.DecimalField(max_digits=10, decimal_places=2)
    image = models.ImageField(upload_to='cars/', blank=True, null=True)
    description = models.TextField()
    is_available = models.BooleanField(default=True)
    mileage = models.IntegerField(default=0)
    category = models.ForeignKey(CarCategory, on_delete=models.SET_NULL, null=True, blank=True)
    color = models.ForeignKey(CarColor, on_delete=models.SET_NULL, null=True, blank=True)
    engine = models.ForeignKey(CarEngine, on_delete=models.SET_NULL, null=True, blank=True)
    features = models.ManyToManyField(CarFeature, blank=True)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    def __str__(self):
        return f"{self.brand} {self.model} ({self.year})"

    class Meta:
        ordering = ['-created_at']

class CarGallery(models.Model):
    car = models.ForeignKey(Car, on_delete=models.CASCADE, related_name='gallery')
    images = models.ImageField(upload_to='car_gallery/')
    created_at = models.DateTimeField(auto_now_add=True)

    def __str__(self):
        return f"Gallery for {self.car}"

    class Meta:
        verbose_name_plural = "Car galleries"
""")

            add_paragraph(doc, "bookings/models.py")
            add_code_block(doc, """from django.db import models
from django.contrib.auth.models import User
from cars.models import Car

class Booking(models.Model):
    STATUS_CHOICES = [
        ('pending', 'Pending'),
        ('confirmed', 'Confirmed'),
        ('cancelled', 'Cancelled'),
        ('completed', 'Completed'),
    ]
    
    user = models.ForeignKey(User, on_delete=models.CASCADE)
    car = models.ForeignKey(Car, on_delete=models.CASCADE)
    start_date = models.DateField()
    end_date = models.DateField()
    total_price = models.DecimalField(max_digits=10, decimal_places=2)
    status = models.CharField(max_length=20, choices=STATUS_CHOICES, default='pending')
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    
    def __str__(self):
        return f"Booking {self.id} - {self.user.username} - {self.car}"
    
    class Meta:
        ordering = ['-created_at']
""")

            add_heading(doc, "11.2 Django Views", 2)
            add_paragraph(doc, "cars/views.py")
            add_code_block(doc, """from rest_framework import viewsets, permissions, filters
from rest_framework.decorators import action
from rest_framework.response import Response
from django_filters.rest_framework import DjangoFilterBackend
from .models import Car, CarCategory, CarColor, CarEngine, CarFeature, CarGallery
from .serializers import CarSerializer, CarCategorySerializer, CarColorSerializer, CarEngineSerializer, CarFeatureSerializer, CarGallerySerializer
from .filters import CarFilter

class CarViewSet(viewsets.ModelViewSet):
    queryset = Car.objects.all()
    serializer_class = CarSerializer
    permission_classes = [permissions.IsAuthenticatedOrReadOnly]
    filter_backends = [DjangoFilterBackend, filters.SearchFilter, filters.OrderingFilter]
    filterset_class = CarFilter
    search_fields = ['brand', 'model', 'description']
    ordering_fields = ['price_per_day', 'year', 'created_at']
    ordering = ['-created_at']

    @action(detail=True, methods=['post'])
    def toggle_availability(self, request, pk=None):
        car = self.get_object()
        car.is_available = not car.is_available
        car.save()
        return Response({'status': 'success', 'is_available': car.is_available})

    @action(detail=False, methods=['get'])
    def available_cars(self, request):
        available_cars = Car.objects.filter(is_available=True)
        serializer = self.get_serializer(available_cars, many=True)
        return Response(serializer.data)

class CarCategoryViewSet(viewsets.ModelViewSet):
    queryset = CarCategory.objects.all()
    serializer_class = CarCategorySerializer
    permission_classes = [permissions.IsAuthenticatedOrReadOnly]

class CarColorViewSet(viewsets.ModelViewSet):
    queryset = CarColor.objects.all()
    serializer_class = CarColorSerializer
    permission_classes = [permissions.IsAuthenticatedOrReadOnly]

class CarEngineViewSet(viewsets.ModelViewSet):
    queryset = CarEngine.objects.all()
    serializer_class = CarEngineSerializer
    permission_classes = [permissions.IsAuthenticatedOrReadOnly]

class CarFeatureViewSet(viewsets.ModelViewSet):
    queryset = CarFeature.objects.all()
    serializer_class = CarFeatureSerializer
    permission_classes = [permissions.IsAuthenticatedOrReadOnly]

class CarGalleryViewSet(viewsets.ModelViewSet):
    queryset = CarGallery.objects.all()
    serializer_class = CarGallerySerializer
    permission_classes = [permissions.IsAuthenticatedOrReadOnly]
""")

            add_paragraph(doc, "bookings/views.py")
            add_code_block(doc, """from rest_framework import viewsets, permissions, status
from rest_framework.decorators import action
from rest_framework.response import Response
from django.utils import timezone
from datetime import datetime
from .models import Booking
from .serializers import BookingSerializer
from .permissions import IsOwnerOrReadOnly

class BookingViewSet(viewsets.ModelViewSet):
    serializer_class = BookingSerializer
    permission_classes = [permissions.IsAuthenticated, IsOwnerOrReadOnly]

    def get_queryset(self):
        if self.request.user.is_staff:
            return Booking.objects.all()
        return Booking.objects.filter(user=self.request.user)

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)

    @action(detail=True, methods=['post'])
    def cancel_booking(self, request, pk=None):
        booking = self.get_object()
        if booking.status == 'confirmed':
            booking.status = 'cancelled'
            booking.save()
            return Response({'status': 'Booking cancelled successfully'})
        return Response({'error': 'Cannot cancel this booking'}, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=True, methods=['post'])
    def confirm_booking(self, request, pk=None):
        booking = self.get_object()
        if booking.status == 'pending':
            booking.status = 'confirmed'
            booking.save()
            return Response({'status': 'Booking confirmed successfully'})
        return Response({'error': 'Cannot confirm this booking'}, status=status.HTTP_400_BAD_REQUEST)
""")

            add_heading(doc, "11.3 React Components", 2)
            add_paragraph(doc, "frontend/src/App.tsx")
            add_code_block(doc, """import React from 'react';
import { BrowserRouter as Router, Routes, Route } from 'react-router-dom';
import { AuthProvider } from './contexts/AuthContext';
import { MaintenanceProvider } from './contexts/MaintenanceContext';
import Navbar from './components/Navbar';
import ProtectedRoute from './components/ProtectedRoute';
import Home from './pages/Home';
import Login from './pages/Login';
import Register from './pages/Register';
import Cars from './pages/Cars';
import CarDetails from './pages/CarDetails';
import Bookings from './pages/Bookings';
import BookingPage from './pages/BookingPage';
import Profile from './pages/Profile';
import Maintenance from './pages/Maintenance';
import AdminDashboard from './pages/Admin/AdminDashboard';
import AdminBookings from './pages/Admin/AdminBookings';
import CarManagement from './pages/Admin/CarManagement';
import AdminSettings from './pages/Admin/AdminSettings';
import './App.css';

function App() {
  return (
    <AuthProvider>
      <MaintenanceProvider>
        <Router>
          <div className="min-h-screen bg-gray-50">
            <Navbar />
            <main className="container mx-auto px-4 py-8">
              <Routes>
                <Route path="/" element={<Home />} />
                <Route path="/login" element={<Login />} />
                <Route path="/register" element={<Register />} />
                <Route path="/cars" element={<Cars />} />
                <Route path="/cars/:id" element={<CarDetails />} />
                <Route path="/bookings" element={<ProtectedRoute><Bookings /></ProtectedRoute>} />
                <Route path="/bookings/new/:carId" element={<ProtectedRoute><BookingPage /></ProtectedRoute>} />
                <Route path="/profile" element={<ProtectedRoute><Profile /></ProtectedRoute>} />
                <Route path="/maintenance" element={<ProtectedRoute><Maintenance /></ProtectedRoute>} />
                <Route path="/admin" element={<ProtectedRoute><AdminDashboard /></ProtectedRoute>} />
                <Route path="/admin/bookings" element={<ProtectedRoute><AdminBookings /></ProtectedRoute>} />
                <Route path="/admin/cars" element={<ProtectedRoute><CarManagement /></ProtectedRoute>} />
                <Route path="/admin/settings" element={<ProtectedRoute><AdminSettings /></ProtectedRoute>} />
              </Routes>
            </main>
          </div>
        </Router>
      </MaintenanceProvider>
    </AuthProvider>
  );
}

export default App;
""")

            add_paragraph(doc, "frontend/src/pages/Cars.tsx")
            add_code_block(doc, """import React, { useState, useEffect } from 'react';
import { Link } from 'react-router-dom';
import { useAuth } from '../contexts/AuthContext';
import { Car } from '../types';

const Cars: React.FC = () => {
  const [cars, setCars] = useState<Car[]>([]);
  const [loading, setLoading] = useState(true);
  const [error, setError] = useState<string | null>(null);
  const [filters, setFilters] = useState({
    brand: '',
    minPrice: '',
    maxPrice: '',
    seats: '',
    year: ''
  });
  const { token } = useAuth();

  useEffect(() => {
    fetchCars();
  }, []);

  const fetchCars = async () => {
    try {
      const response = await fetch('http://localhost:8000/api/cars/', {
        headers: {
          'Authorization': `Bearer ${token}`,
          'Content-Type': 'application/json',
        },
      });
      
      if (!response.ok) {
        throw new Error('Failed to fetch cars');
      }
      
      const data = await response.json();
      setCars(data);
    } catch (err) {
      setError(err instanceof Error ? err.message : 'An error occurred');
    } finally {
      setLoading(false);
    }
  };

  const handleFilterChange = (e: React.ChangeEvent<HTMLInputElement | HTMLSelectElement>) => {
    const { name, value } = e.target;
    setFilters(prev => ({
      ...prev,
      [name]: value
    }));
  };

  const filteredCars = cars.filter(car => {
    if (filters.brand && !car.brand.toLowerCase().includes(filters.brand.toLowerCase())) return false;
    if (filters.minPrice && car.price_per_day < parseFloat(filters.minPrice)) return false;
    if (filters.maxPrice && car.price_per_day > parseFloat(filters.maxPrice)) return false;
    if (filters.seats && car.seats !== parseInt(filters.seats)) return false;
    if (filters.year && car.year !== parseInt(filters.year)) return false;
    return true;
  });

  if (loading) {
    return (
      <div className="flex justify-center items-center h-64">
        <div className="animate-spin rounded-full h-32 w-32 border-b-2 border-blue-600"></div>
      </div>
    );
  }

  if (error) {
    return (
      <div className="text-center py-8">
        <div className="text-red-600 text-xl mb-4">Error: {error}</div>
        <button 
          onClick={fetchCars}
          className="bg-blue-600 text-white px-4 py-2 rounded hover:bg-blue-700"
        >
          Retry
        </button>
      </div>
    );
  }

  return (
    <div className="max-w-7xl mx-auto">
      <div className="mb-8">
        <h1 className="text-3xl font-bold text-gray-900 mb-4">Available Cars</h1>
        <p className="text-gray-600">Discover our premium collection of vehicles</p>
      </div>

      {/* Filters */}
      <div className="bg-white p-6 rounded-lg shadow-md mb-8">
        <h2 className="text-xl font-semibold mb-4">Filters</h2>
        <div className="grid grid-cols-1 md:grid-cols-2 lg:grid-cols-5 gap-4">
          <input
            type="text"
            name="brand"
            placeholder="Brand"
            value={filters.brand}
            onChange={handleFilterChange}
            className="border border-gray-300 rounded px-3 py-2 focus:outline-none focus:ring-2 focus:ring-blue-500"
          />
          <input
            type="number"
            name="minPrice"
            placeholder="Min Price"
            value={filters.minPrice}
            onChange={handleFilterChange}
            className="border border-gray-300 rounded px-3 py-2 focus:outline-none focus:ring-2 focus:ring-blue-500"
          />
          <input
            type="number"
            name="maxPrice"
            placeholder="Max Price"
            value={filters.maxPrice}
            onChange={handleFilterChange}
            className="border border-gray-300 rounded px-3 py-2 focus:outline-none focus:ring-2 focus:ring-blue-500"
          />
          <select
            name="seats"
            value={filters.seats}
            onChange={handleFilterChange}
            className="border border-gray-300 rounded px-3 py-2 focus:outline-none focus:ring-2 focus:ring-blue-500"
          >
            <option value="">All Seats</option>
            <option value="2">2 Seats</option>
            <option value="4">4 Seats</option>
            <option value="5">5 Seats</option>
            <option value="7">7 Seats</option>
          </select>
          <input
            type="number"
            name="year"
            placeholder="Year"
            value={filters.year}
            onChange={handleFilterChange}
            className="border border-gray-300 rounded px-3 py-2 focus:outline-none focus:ring-2 focus:ring-blue-500"
          />
        </div>
      </div>

      {/* Cars Grid */}
      <div className="grid grid-cols-1 md:grid-cols-2 lg:grid-cols-3 xl:grid-cols-4 gap-6">
        {filteredCars.map((car) => (
          <div key={car.id} className="bg-white rounded-lg shadow-md overflow-hidden hover:shadow-lg transition-shadow">
            <div className="relative">
              {car.image ? (
                <img
                  src={car.image}
                  alt={`${car.brand} ${car.model}`}
                  className="w-full h-48 object-cover"
                />
              ) : (
                <div className="w-full h-48 bg-gray-200 flex items-center justify-center">
                  <span className="text-gray-500">No Image</span>
                </div>
              )}
              <div className="absolute top-2 right-2">
                <span className={`px-2 py-1 rounded-full text-xs font-semibold ${
                  car.is_available 
                    ? 'bg-green-100 text-green-800' 
                    : 'bg-red-100 text-red-800'
                }`}>
                  {car.is_available ? 'Available' : 'Unavailable'}
                </span>
              </div>
            </div>
            
            <div className="p-4">
              <h3 className="text-lg font-semibold text-gray-900 mb-2">
                {car.brand} {car.model}
              </h3>
              <p className="text-gray-600 text-sm mb-2">Year: {car.year}</p>
              <p className="text-gray-600 text-sm mb-3">Seats: {car.seats}</p>
              <div className="flex justify-between items-center">
                <span className="text-xl font-bold text-blue-600">
                  ${car.price_per_day}/day
                </span>
                <Link
                  to={`/cars/${car.id}`}
                  className="bg-blue-600 text-white px-4 py-2 rounded hover:bg-blue-700 transition-colors"
                >
                  View Details
                </Link>
              </div>
            </div>
          </div>
        ))}
      </div>

      {filteredCars.length === 0 && (
        <div className="text-center py-12">
          <p className="text-gray-500 text-lg">No cars found matching your criteria.</p>
        </div>
      )}
    </div>
  );
};

export default Cars;
""")

            add_heading(doc, "11.4 Configuration Files", 2)
            add_paragraph(doc, "car_booking/settings.py")
            add_code_block(doc, """import os
from pathlib import Path
from datetime import timedelta

# Build paths inside the project like this: BASE_DIR / 'subdir'.
BASE_DIR = Path(__file__).resolve().parent.parent

# SECURITY WARNING: keep the secret key used in production secret!
SECRET_KEY = 'django-insecure-your-secret-key-here'

# SECURITY WARNING: don't run with debug turned on in production!
DEBUG = True

ALLOWED_HOSTS = []

# Application definition
INSTALLED_APPS = [
    'django.contrib.admin',
    'django.contrib.auth',
    'django.contrib.contenttypes',
    'django.contrib.sessions',
    'django.contrib.messages',
    'django.contrib.staticfiles',
    'rest_framework',
    'rest_framework_simplejwt',
    'corsheaders',
    'django_filters',
    'users',
    'cars',
    'bookings',
]

MIDDLEWARE = [
    'django.middleware.security.SecurityMiddleware',
    'django.contrib.sessions.middleware.SessionMiddleware',
    'corsheaders.middleware.CorsMiddleware',
    'django.middleware.common.CommonMiddleware',
    'django.middleware.csrf.CsrfViewMiddleware',
    'django.contrib.auth.middleware.AuthenticationMiddleware',
    'django.contrib.messages.middleware.MessageMiddleware',
    'django.middleware.clickjacking.XFrameOptionsMiddleware',
]

ROOT_URLCONF = 'car_booking.urls'

TEMPLATES = [
    {
        'BACKEND': 'django.template.backends.django.DjangoTemplates',
        'DIRS': [],
        'APP_DIRS': True,
        'OPTIONS': {
            'context_processors': [
                'django.template.context_processors.debug',
                'django.template.context_processors.request',
                'django.contrib.auth.context_processors.auth',
                'django.contrib.messages.context_processors.messages',
            ],
        },
    },
]

WSGI_APPLICATION = 'car_booking.wsgi.application'

# Database
DATABASES = {
    'default': {
        'ENGINE': 'django.db.backends.sqlite3',
        'NAME': BASE_DIR / 'db.sqlite3',
    }
}

# Password validation
AUTH_PASSWORD_VALIDATORS = [
    {
        'NAME': 'django.contrib.auth.password_validation.UserAttributeSimilarityValidator',
    },
    {
        'NAME': 'django.contrib.auth.password_validation.MinimumLengthValidator',
    },
    {
        'NAME': 'django.contrib.auth.password_validation.CommonPasswordValidator',
    },
    {
        'NAME': 'django.contrib.auth.password_validation.NumericPasswordValidator',
    },
]

# Internationalization
LANGUAGE_CODE = 'en-us'
TIME_ZONE = 'UTC'
USE_I18N = True
USE_TZ = True

# Static files (CSS, JavaScript, Images)
STATIC_URL = 'static/'
STATIC_ROOT = os.path.join(BASE_DIR, 'staticfiles')
MEDIA_URL = '/media/'
MEDIA_ROOT = os.path.join(BASE_DIR, 'media')

# Default primary key field type
DEFAULT_AUTO_FIELD = 'django.db.models.BigAutoField'

# REST Framework settings
REST_FRAMEWORK = {
    'DEFAULT_AUTHENTICATION_CLASSES': (
        'rest_framework_simplejwt.authentication.JWTAuthentication',
    ),
    'DEFAULT_PERMISSION_CLASSES': (
        'rest_framework.permissions.IsAuthenticated',
    ),
    'DEFAULT_FILTER_BACKENDS': (
        'django_filters.rest_framework.DjangoFilterBackend',
        'rest_framework.filters.SearchFilter',
        'rest_framework.filters.OrderingFilter',
    ),
    'DEFAULT_PAGINATION_CLASS': 'rest_framework.pagination.PageNumberPagination',
    'PAGE_SIZE': 10,
}

# JWT settings
SIMPLE_JWT = {
    'ACCESS_TOKEN_LIFETIME': timedelta(minutes=60),
    'REFRESH_TOKEN_LIFETIME': timedelta(days=1),
    'ROTATE_REFRESH_TOKENS': False,
    'BLACKLIST_AFTER_ROTATION': True,
    'UPDATE_LAST_LOGIN': False,
    'ALGORITHM': 'HS256',
    'SIGNING_KEY': SECRET_KEY,
    'VERIFYING_KEY': None,
    'AUDIENCE': None,
    'ISSUER': None,
    'AUTH_HEADER_TYPES': ('Bearer',),
    'AUTH_HEADER_NAME': 'HTTP_AUTHORIZATION',
    'USER_ID_FIELD': 'id',
    'USER_ID_CLAIM': 'user_id',
    'AUTH_TOKEN_CLASSES': ('rest_framework_simplejwt.tokens.AccessToken',),
    'TOKEN_TYPE_CLAIM': 'token_type',
    'JTI_CLAIM': 'jti',
    'SLIDING_TOKEN_REFRESH_EXP_CLAIM': 'refresh_exp',
    'SLIDING_TOKEN_LIFETIME': timedelta(minutes=5),
    'SLIDING_TOKEN_REFRESH_LIFETIME': timedelta(days=1),
}

# CORS settings
CORS_ALLOWED_ORIGINS = [
    "http://localhost:3000",
    "http://127.0.0.1:3000",
]

CORS_ALLOW_CREDENTIALS = True

# Custom user model
AUTH_USER_MODEL = 'auth.User'
""")

            add_paragraph(doc, "frontend/vite.config.ts")
            add_code_block(doc, """import { defineConfig } from 'vite'
import react from '@vitejs/plugin-react'

// https://vitejs.dev/config/
export default defineConfig({
  plugins: [react()],
  server: {
    port: 3000,
    host: true,
  },
  build: {
    outDir: 'dist',
    sourcemap: true,
  },
  define: {
    global: 'globalThis',
  },
})
""")

            add_paragraph(doc, "frontend/tailwind.config.js")
            add_code_block(doc, """/** @type {import('tailwindcss').Config} */
export default {
  content: [
    "./index.html",
    "./src/**/*.{js,ts,jsx,tsx}",
  ],
  theme: {
    extend: {
      colors: {
        primary: {
          50: '#eff6ff',
          100: '#dbeafe',
          200: '#bfdbfe',
          300: '#93c5fd',
          400: '#60a5fa',
          500: '#3b82f6',
          600: '#2563eb',
          700: '#1d4ed8',
          800: '#1e40af',
          900: '#1e3a8a',
          950: '#172554',
        },
      },
      fontFamily: {
        sans: ['Inter', 'system-ui', 'sans-serif'],
      },
      animation: {
        'fade-in': 'fadeIn 0.5s ease-in-out',
        'slide-up': 'slideUp 0.3s ease-out',
      },
      keyframes: {
        fadeIn: {
          '0%': { opacity: '0' },
          '100%': { opacity: '1' },
        },
        slideUp: {
          '0%': { transform: 'translateY(10px)', opacity: '0' },
          '100%': { transform: 'translateY(0)', opacity: '1' },
        },
      },
    },
  },
  plugins: [],
}
""")

            break
    
    if not section_10_found:
        print("Section 10 heading not found")

def main():
    doc = docx.Document()
    doc.add_heading('GrandCruise: Full-Stack Premium Car Booking & Rental Management Platform', 0)
    
    # 1. INTRODUCTION
    add_heading(doc, '1. INTRODUCTION', 1)
    add_paragraph(doc, """
The landscape of personal and business mobility has undergone a significant transformation in recent years, driven by rapid advancements in digital technology and changing consumer expectations. In the premium car rental sector, customers now demand not only access to high-end vehicles but also a seamless, intuitive, and secure digital experience that matches the luxury of the cars themselves. Traditional rental models, often reliant on manual processes and outdated systems, are increasingly unable to meet these evolving needs.

GrandCruise emerges as a response to these challenges, offering a full-stack, premium car booking and rental management platform that redefines the standards of luxury mobility services. The platform is meticulously designed to bridge the gap between elite car rental providers and discerning clients, providing a digital ecosystem where convenience, security, and sophistication converge. GrandCruise leverages state-of-the-art web technologies to deliver a robust, scalable, and user-friendly solution that caters to both customers and administrators.

For customers, GrandCruise offers an unparalleled experience: from browsing a curated selection of luxury vehicles—such as BMW, Mercedes-Benz, Audi, and Porsche—to making real-time bookings, managing reservations, and accessing personalized services. The platform's intuitive interface, real-time availability, and secure payment processing ensure that every interaction is smooth and trustworthy.

For businesses, GrandCruise provides powerful tools for fleet management, analytics, and customer relationship management. Administrators can efficiently oversee operations, optimize fleet utilization, and gain actionable insights through comprehensive dashboards and reporting features. The platform's modular architecture and integration capabilities allow for easy adaptation to business growth and technological advancements.

By setting a new benchmark for digital excellence in the premium car rental industry, GrandCruise empowers both providers and customers to embrace the future of luxury mobility with confidence and ease.
""")

    # 2. EXISTING SYSTEM
    add_heading(doc, '2. EXISTING SYSTEM', 1)
    add_paragraph(doc, """
The existing landscape of premium car rental systems is characterized by a reliance on legacy software, fragmented processes, and limited digital integration. Many providers still operate with systems that were originally designed for standard vehicle rentals, lacking the sophistication and flexibility required for high-end, luxury fleets. These platforms often fail to deliver the seamless, on-demand experience that modern customers expect, especially in an era where digital convenience is paramount.

Customers seeking to rent premium vehicles frequently encounter outdated websites or mobile apps, slow response times, and a lack of transparency regarding vehicle availability and pricing. The absence of real-time data synchronization can lead to double bookings, missed opportunities, and customer dissatisfaction. On the administrative side, managers and staff are often burdened with manual data entry, paper-based contracts, and disjointed communication channels, making it difficult to maintain operational efficiency and deliver a truly premium service.

Furthermore, existing systems rarely offer advanced features such as dynamic pricing, personalized recommendations, or integrated analytics, which are increasingly important for both customer engagement and business optimization. Security and compliance with modern data protection standards are also common concerns, as many legacy platforms were not built with today's cybersecurity threats in mind.
""")
    
    # 2.1. Disadvantages
    add_heading(doc, '2.1. Disadvantages', 2)
    add_paragraph(doc, '- Manual Processes: Many existing solutions require manual intervention for booking confirmations, vehicle assignments, and record-keeping, leading to errors, inefficiencies, and delays.')
    add_paragraph(doc, '- Limited User Experience: Outdated interfaces, lack of mobile responsiveness, and unintuitive navigation hinder customer engagement and satisfaction.')
    add_paragraph(doc, '- No Real-Time Updates: The absence of real-time vehicle availability and booking status can result in double bookings, missed opportunities, and customer frustration.')
    add_paragraph(doc, '- Poor Integration: Existing systems often lack integration with payment gateways, analytics, CRM tools, and third-party services, resulting in fragmented workflows and data silos.')
    add_paragraph(doc, '- Security and Compliance Risks: Legacy systems may not adhere to modern security standards or data protection regulations, putting sensitive customer and business data at risk of breaches or non-compliance.')
    add_paragraph(doc, '- Lack of Advanced Features: Most current platforms do not support dynamic pricing, personalized recommendations, or advanced analytics, limiting the ability to optimize business operations and enhance customer experience.')

    # 3. PROPOSED SYSTEM
    add_heading(doc, '3. PROPOSED SYSTEM', 1)
    add_paragraph(doc, """
GrandCruise introduces a transformative, full-stack digital solution specifically designed for the premium car rental industry. Recognizing the unique demands of luxury vehicle rentals, GrandCruise is built to deliver a seamless, secure, and highly efficient experience for both customers and business operators. The platform leverages the latest advancements in web technology to automate and streamline every aspect of the rental process, from the initial vehicle search to final booking and fleet management.

At its core, GrandCruise is engineered to provide real-time access to a curated selection of high-end vehicles, ensuring that customers can easily browse, compare, and reserve their preferred models with confidence. The system's robust backend infrastructure guarantees data integrity, security, and rapid response times, while the intuitive frontend interface is optimized for both desktop and mobile users, reflecting the expectations of a modern, tech-savvy clientele.

For administrators and business owners, GrandCruise offers a centralized platform to manage fleet inventory, oversee bookings, and maintain customer relationships. The system is designed to minimize manual intervention, reduce operational overhead, and eliminate common pain points such as double bookings or miscommunication. Built-in tools for secure payment processing, user authentication, and access control further enhance the reliability and trustworthiness of the platform.

By integrating all essential functions into a single, cohesive ecosystem, GrandCruise empowers luxury car rental businesses to scale efficiently, adapt to market trends, and deliver a truly premium service experience. The platform's modular architecture also allows for future enhancements and integrations, ensuring long-term value and flexibility as the industry evolves.
""")

    # 3.1. Proposed Models
    add_heading(doc, '3.1. Proposed Models', 2)
    add_paragraph(doc, '- User Module: Enables customers to register, browse available premium vehicles, make bookings, and manage their profiles.')
    add_paragraph(doc, '- Admin Module: Provides administrators with tools for fleet management, booking oversight, and customer support.')
    add_paragraph(doc, '- Car Management Module: Allows for the addition, modification, and removal of vehicles, including detailed specifications and image galleries.')
    add_paragraph(doc, '- Booking Management Module: Handles real-time booking requests, availability checks, payment processing, and status updates.')

    # 3.2. Advantages
    add_heading(doc, '3.2. Advantages', 2)
    add_paragraph(doc, '- Automated Processes: GrandCruise eliminates manual intervention by automating booking confirmations, vehicle assignments, and record-keeping, reducing errors and operational delays.')
    add_paragraph(doc, '- Superior User Experience: The platform features a modern, intuitive, and mobile-responsive interface, ensuring seamless navigation and high customer satisfaction.')
    add_paragraph(doc, '- Real-Time Data Synchronization: GrandCruise provides instant updates on vehicle availability and booking status, preventing double bookings and missed opportunities.')
    add_paragraph(doc, '- Seamless Integration: The system integrates smoothly with payment gateways, CRM tools, and third-party services, unifying workflows and eliminating data silos.')
    add_paragraph(doc, '- Advanced Security & Compliance: GrandCruise adheres to the latest security standards and data protection regulations, safeguarding sensitive customer and business information.')
    add_paragraph(doc, '- Advanced Features: The platform supports dynamic pricing, personalized recommendations, and analytics, empowering businesses to optimize operations and enhance the customer experience.')

    # 4. SYSTEM ARCHITECTURE
    add_heading(doc, '4. SYSTEM ARCHITECTURE', 1)
    add_paragraph(doc, """
GrandCruise employs a modular, layered architecture designed for scalability and maintainability. The frontend layer, built with React, provides an intuitive user interface for both customers and administrators, handling navigation and state management through contexts and shared components. The backend layer, powered by Django and Django REST Framework, exposes RESTful APIs that encapsulate all business logic, including authentication, user management, car management, and booking operations. Data persistence is handled by a relational database (PostgreSQL or MySQL) that stores users, cars, bookings, and system settings, while media assets are managed through dedicated storage solutions. The architecture implements secure authentication and role-based access control throughout the stack, ensuring data protection and user privacy. This design supports both cloud and on-premises deployment, enabling high availability and seamless scaling as business requirements evolve.
""")
    add_paragraph(doc, 'GrandCruise adopts a modular, layered architecture:')
    add_paragraph(doc, '- Frontend: Built with React, providing a dynamic, responsive, and intuitive user interface for both customers and administrators.')
    add_paragraph(doc, '- Backend: Powered by Django, offering robust RESTful APIs, business logic, and secure data management.')
    add_paragraph(doc, '- Database: Utilizes PostgreSQL (or similar RDBMS) for reliable, scalable data storage.')
    add_paragraph(doc, '- Media Storage: Handles car images and documents via a dedicated media server or cloud storage.')
    add_paragraph(doc, '- Authentication: Implements secure user authentication and role-based access control.')
    add_paragraph(doc, '- Deployment: Supports cloud or on-premises deployment, with scalability and high availability.')
    add_paragraph(doc, '[Optional: Add a system architecture diagram if needed.]')

    # 5. SOFTWARE AND HARDWARE REQUIREMENTS
    add_heading(doc, '5. SOFTWARE AND HARDWARE REQUIREMENTS', 1)
    add_paragraph(doc, 'Software Requirements:')
    add_paragraph(doc, '- Operating System: Windows, Linux, or macOS')
    add_paragraph(doc, '- Backend Framework: Django (Python 3.8+)')
    add_paragraph(doc, '- Frontend Framework: React (Node.js 16+)')
    add_paragraph(doc, '- Database: PostgreSQL (or MySQL/SQLite for development)')
    add_paragraph(doc, '- Web Server: Nginx or Apache (for production)')
    add_paragraph(doc, '- Other Tools: Docker (optional), Git, Postman, VS Code')
    add_paragraph(doc, 'Hardware Requirements:')
    add_paragraph(doc, '- Development Machine:')
    add_paragraph(doc, '  - Processor: Intel i5 or higher (or equivalent)')
    add_paragraph(doc, '  - RAM: 8 GB minimum (16 GB recommended)')
    add_paragraph(doc, '  - Storage: 50 GB free disk space')
    add_paragraph(doc, '- Production Server:')
    add_paragraph(doc, '  - Processor: Quad-core CPU or higher')
    add_paragraph(doc, '  - RAM: 16 GB or higher')
    add_paragraph(doc, '  - Storage: SSD with 100 GB or more (depending on media storage needs)')
    add_paragraph(doc, '  - Network: High-speed internet connection')

    # 6. SYSTEM IMPLEMENTATION
    add_heading(doc, '6. SYSTEM IMPLEMENTATION', 1)
    add_paragraph(doc, """
The implementation of GrandCruise follows a comprehensive modular approach with clear separation of concerns between frontend and backend components, ensuring maintainability, scalability, and robust performance. The system architecture is built using Django 5.0.2 with Django REST Framework for the backend API, providing a robust foundation for business logic, data management, and security. The frontend interface is developed using React 18 with TypeScript, offering a dynamic, responsive, and intuitive user experience. MySQL serves as the primary database, providing reliable data persistence with ACID compliance and optimized query performance. The implementation incorporates modern development practices including RESTful API design, token-based authentication, role-based access control, and responsive design principles. The development methodology emphasizes code reusability, component-based architecture, and comprehensive error handling. The system integrates various third-party libraries and tools such as React Query for state management, Tailwind CSS for styling, and django-filter for advanced data filtering capabilities. The following sub-sections detail the key implementation aspects of the platform, covering database design, API development, frontend architecture, security implementation, and administrative functionality.
""")

    # 6.1. Django Models & Database Design
    add_heading(doc, '6.1. Django Models & Database Design', 2)
    add_paragraph(doc, """
The database schema is implemented using Django's ORM with four main models: User, Car, Booking, and SystemSettings. The User model extends Django's built-in User model with a UserProfile that includes phone number and address fields. The Car model stores comprehensive vehicle information including brand, model, year, price, specifications, and availability status. The Booking model manages reservation data with status tracking (pending, confirmed, cancelled, completed) and automatic price calculation based on rental duration. The SystemSettings model handles global configuration like maintenance mode and booking policies. All models include created_at and updated_at timestamps for audit trails, and the database uses MySQL with proper indexing for optimal query performance.
""")

    # 6.2. REST API Implementation
    add_heading(doc, '6.2. REST API Implementation', 2)
    add_paragraph(doc, """
The REST API is built using Django REST Framework with ViewSets for each model. The API includes endpoints for user management (/api/users/), car catalog (/api/cars/), booking operations (/api/bookings/), and system settings (/api/admin/settings/). Authentication is handled through token-based authentication with support for session and basic authentication. The API implements proper permissions with IsAuthenticated as default and custom permissions like IsOwnerOrAdmin for booking operations. Serializers handle data validation and transformation, with nested serialization for related objects. The API supports filtering, searching, and ordering capabilities using django-filter and DRF's built-in features.
""")

    # 6.3. Frontend React Components
    add_heading(doc, '6.3. Frontend React Components', 2)
    add_paragraph(doc, """
The frontend is built with React and TypeScript, using React Router for navigation and React Query for state management. The application structure includes pages for home, car browsing, booking, user authentication, and admin functions. Context providers (AuthContext and MaintenanceContext) manage global state for user authentication and system maintenance mode. Shared components include navigation bars, cards, and form elements styled with Tailwind CSS. The frontend communicates with the backend API using Axios for HTTP requests, with proper error handling and loading states. Protected routes ensure only authenticated users can access booking and profile features.
""")

    # 6.4. Authentication & Authorization
    add_heading(doc, '6.4. Authentication & Authorization', 2)
    add_paragraph(doc, """
Authentication is implemented using Django's built-in authentication system with token-based authentication for API access. User registration and login are handled through the /api/auth/login/ endpoint, which returns authentication tokens. The frontend stores tokens in localStorage and includes them in API request headers. Role-based access control is enforced through Django permissions, with staff users having access to admin functions. Custom permissions ensure users can only access their own bookings while admins can view all data. The system includes maintenance mode functionality that can restrict access to certain features during system updates.
""")

    # 6.5. Admin Panel Implementation
    add_heading(doc, '6.5. Admin Panel Implementation', 2)
    add_paragraph(doc, """
The admin panel provides comprehensive management capabilities through both Django's built-in admin interface and custom React components. Django admin offers CRUD operations for all models with proper permissions. Custom admin pages built in React include dashboard views, car management with image uploads, booking oversight, and system settings configuration. The admin interface includes features for managing user accounts, car inventory with detailed specifications, booking status updates, and system-wide settings like maintenance mode and booking policies. All admin operations are protected by authentication and authorization checks.
""")

    # 7. PYTHON, REACT, AND TAILWIND CSS TECHNOLOGIES
    add_heading(doc, '7. PYTHON, REACT, AND TAILWIND CSS TECHNOLOGIES', 1)
    
    # 7.1. Python Programming Language
    add_heading(doc, '7.1. Python Programming Language', 2)
    
    # 7.1.1. History of Python
    add_heading(doc, '7.1.1. History of Python', 3)
    add_paragraph(doc, """
Python was created by Guido van Rossum and first released in 1991. It was designed to emphasize code readability and simplicity, making it accessible to both beginners and experienced developers. Python's development began in the late 1980s when Guido was working at the Centrum Wiskunde & Informatica (CWI) in the Netherlands. He started the project as a successor to the ABC language, which was designed for teaching and prototyping but had limitations that Guido aimed to address.
""")
    add_paragraph(doc, """
During Christmas of 1989, Guido began implementing Python as a hobby project to keep himself occupied during the holidays. His goal was to create an easy-to-use language with a broad set of utilities and extensibility. The first version of Python (version 0.9.0) was released to the public in February 1991. This early version already included many of Python's core features, such as exception handling, functions, and the core data types of list, dict, str, and more.
""")
    add_paragraph(doc, """
Python 1.0 was released in 1994, marking Python's first major public release. This version included new features like functional programming tools (lambda, map, filter, and reduce) and built-in support for complex numbers. Python 2.0 was released in 2000, introducing many features that improved the language's functionality and usability, such as list comprehensions, a full garbage collector, and support for Unicode.
""")
    add_paragraph(doc, """
Python 3.0 was released in 2008, representing a significant update that was not backward compatible with previous versions. Python 3 was designed to fix inherent design flaws in the language and remove redundant constructs. Some major changes included the print function, new syntax for integer division, and improvements in Unicode handling. Python continues to evolve with new releases, bringing more features and improvements while maintaining its position as one of the most popular programming languages worldwide.
""")
    
    # 7.1.2. Introduction to Python
    add_heading(doc, '7.1.2. Introduction to Python', 3)
    add_paragraph(doc, """
Python is a high-level, interpreted programming language renowned for its simplicity, readability, and versatility. Created by Guido van Rossum and first released in 1991, Python has evolved into one of the most widely-used programming languages in the world, embraced by both beginners and experienced developers. Its design philosophy emphasizes code readability, simplicity, and the ability to express concepts in fewer lines of code compared to other languages like C++ or Java.
""")
    add_paragraph(doc, """
Python's syntax is clean and easy to read, which reduces the cost of program maintenance. The language supports modules and packages, encouraging program modularity and code reuse. The Python interpreter and the extensive standard library are available in source or binary form without charge for all major platforms and can be freely distributed. Python supports multiple programming paradigms, including procedural, object-oriented, and functional programming.
""")
    
    # 7.1.3. Goal of Python
    add_heading(doc, '7.1.3. Goal of Python', 3)
    add_paragraph(doc, """
The primary goal of Python is to provide a straightforward and efficient way to develop software. Python's design philosophy prioritizes the following principles: Readability - Python syntax is clean and easy to read, making it accessible for beginners and enhancing productivity for experienced developers. Simplicity - Python emphasizes simplicity and clarity, enabling developers to write less code and focus more on solving problems.
""")
    add_paragraph(doc, """
Flexibility - Python is a general-purpose language that can be used for a wide variety of applications, from web development and automation to data analysis and artificial intelligence. Comprehensibility - Python aims to be intuitive, allowing developers to quickly understand and modify code, which is particularly beneficial for collaborative projects. The language promotes rapid development and clean, pragmatic design.
""")
    
    # 7.1.4. Execution Process in Python
    add_heading(doc, '7.1.4. Execution Process in Python', 3)
    add_paragraph(doc, """
Python code is written in plain text files with a .py extension. A Python script can include statements, functions, classes, and modules. Python is an interpreted language, executing code line by line. To run a Python script, you use the Python interpreter through the command line or an integrated development environment (IDE) like PyCharm, VSCode, or Jupyter Notebook.
""")
    add_paragraph(doc, """
When a Python script is run, the interpreter first compiles the code into an intermediate form known as bytecode. This step is automatic and transparent to the user. Bytecode is a lower-level, platform-independent representation of your source code. The compiled bytecode is interpreted by the Python Virtual Machine (PVM). The PVM reads and executes the bytecode instructions, interacting with the underlying operating system to perform tasks such as memory management and I/O operations.
""")
    add_paragraph(doc, """
Python uses automatic memory management, primarily through garbage collection. The garbage collector automatically deallocates memory that is no longer in use, helping to prevent memory leaks and optimize memory usage. Python provides robust error and exception handling mechanisms. Errors can be caught and managed using try, except, else, and finally blocks. Python offers various tools for debugging, such as pdb (Python Debugger) and IDE-specific debuggers.
""")
    
    # 7.1.5. Libraries and Packages in Python
    add_heading(doc, '7.1.5. Libraries and Packages in Python', 3)
    add_paragraph(doc, """
Python's extensive standard library and its vibrant ecosystem of third-party packages make it a powerful tool for a wide range of applications. The standard library includes modules like os (operating system interface), sys (system-specific parameters), datetime (date and time utilities), re (regular expressions), math (mathematical functions), and many more. These built-in modules provide essential functionality without requiring additional installation.
""")
    add_paragraph(doc, """
For web development, Django and Flask are popular frameworks. Django is a high-level framework that encourages rapid development and clean, pragmatic design, while Flask is a micro-framework that provides the basic tools to build a web application. For data science, libraries like Pandas, NumPy, and Matplotlib are widely used. Pandas offers data structures like DataFrame and Series for data manipulation and analysis.
""")
    add_paragraph(doc, """
For machine learning, libraries such as Scikit-learn, TensorFlow, and Keras are essential. Scikit-learn provides simple and efficient tools for data mining and data analysis, including classification, regression, clustering, and dimensionality reduction algorithms. TensorFlow and Keras offer comprehensive tools for building and training machine learning models, including neural networks.
""")
    
    # 7.1.6. Uses of Python
    add_heading(doc, '7.1.6. Uses of Python', 3)
    add_paragraph(doc, """
Python is a versatile programming language widely used in various domains due to its simplicity and readability. Web Development: Python frameworks like Django, Flask, and Pyramid facilitate the development of robust and scalable web applications. These frameworks provide tools for handling HTTP requests, database operations, user authentication, and more.
""")
    add_paragraph(doc, """
Data Science and Analytics: Libraries like Pandas, NumPy, and SciPy are essential for data manipulation, analysis, and scientific computing. Python is the language of choice for many data scientists due to its powerful data processing capabilities and extensive visualization libraries like Matplotlib and Seaborn.
""")
    add_paragraph(doc, """
Machine Learning and Artificial Intelligence: Python's extensive libraries, such as Scikit-learn, TensorFlow, and Keras, support the development and deployment of machine learning models and AI applications. The language's simplicity and the availability of pre-built algorithms make it ideal for AI development.
""")
    add_paragraph(doc, """
Automation and Scripting: Python's simplicity makes it ideal for automating repetitive tasks and scripting, improving productivity and efficiency. The language can interact with operating systems, manipulate files, and control other applications through various APIs.
""")
    
    # 7.1.7. Advantages of Python
    add_heading(doc, '7.1.7. Advantages of Python', 3)
    add_paragraph(doc, """
Ease of Learning: Python's syntax is designed to be easy to understand and write, making it an excellent choice for beginners. The language uses indentation to define code blocks, which enforces good coding practices and makes code more readable.
""")
    add_paragraph(doc, """
Extensive Support Libraries: Python has a rich set of libraries and frameworks that facilitate development in various domains. The Python Package Index (PyPI) contains thousands of packages that can be easily installed using pip, Python's package installer.
""")
    add_paragraph(doc, """
Community and Ecosystem: Python has a large, active community that contributes to its vast ecosystem of open-source libraries and tools. The community provides extensive documentation, tutorials, and support through forums, conferences, and online resources.
""")
    add_paragraph(doc, """
Cross-Platform Compatibility: Python runs on multiple platforms, including Windows, macOS, Linux, and more. This portability makes it easy to develop applications that can run on different operating systems without modification.
""")
    add_paragraph(doc, """
Integration Capabilities: Python can easily integrate with other languages and technologies, such as C/C++, Java, and .NET. This flexibility allows developers to use Python for specific parts of a project while leveraging other technologies where appropriate.
""")
    
    # 7.2. React JavaScript Library
    add_heading(doc, '7.2. React JavaScript Library', 2)
    
    # 7.2.1. History of React
    add_heading(doc, '7.2.1. History of React', 3)
    add_paragraph(doc, """
React was created by Jordan Walke, a software engineer at Facebook, and was first deployed on Facebook's News Feed in 2011 and later on Instagram in 2012. It was developed to address the challenges of building complex user interfaces with data that changes over time. React was initially developed as an internal tool at Facebook to solve the problem of building large applications with data that changes frequently.
""")
    add_paragraph(doc, """
React was first released to the public in May 2013 at the JSConf US conference. The library was open-sourced and quickly gained popularity among developers due to its innovative approach to building user interfaces. React's component-based architecture and virtual DOM concept revolutionized how developers think about building web applications.
""")
    add_paragraph(doc, """
Over the years, React has evolved significantly with major releases introducing new features and improvements. React 16, released in 2017, introduced features like Fragments, Error Boundaries, and Portals. React 18, released in 2022, brought concurrent features, automatic batching, and improved server-side rendering capabilities.
""")
    
    # 7.2.2. Introduction to React
    add_heading(doc, '7.2.2. Introduction to React', 3)
    add_paragraph(doc, """
React is a JavaScript library for building user interfaces, particularly single-page applications. It was developed by Facebook and is used to create interactive UIs efficiently. React allows developers to create large web applications that can change data without reloading the page. Its main goal is to be fast, scalable, and simple.
""")
    add_paragraph(doc, """
React uses a component-based architecture, where the UI is broken down into reusable components. Each component manages its own state and can be composed to create complex user interfaces. React uses a virtual DOM (Document Object Model) to improve performance by minimizing direct manipulation of the actual DOM.
""")
    add_paragraph(doc, """
React follows a declarative programming paradigm, meaning developers describe what they want the UI to look like, and React handles the DOM updates to match that description. This approach makes the code more predictable and easier to debug compared to imperative approaches.
""")
    
    # 7.2.3. Key Features of React
    add_heading(doc, '7.2.3. Key Features of React', 3)
    add_paragraph(doc, """
Virtual DOM: React uses a virtual DOM to improve performance. When the state of a component changes, React creates a new virtual DOM tree and compares it with the previous one. Only the differences are applied to the actual DOM, minimizing expensive DOM operations and improving performance.
""")
    add_paragraph(doc, """
Component-Based Architecture: React applications are built using components, which are reusable pieces of UI. Components can be functional (using hooks) or class-based. They can accept props (properties) and manage their own state, making them highly modular and reusable.
""")
    add_paragraph(doc, """
JSX: React uses JSX (JavaScript XML), a syntax extension that allows developers to write HTML-like code in JavaScript. JSX makes it easier to visualize the structure of components and their relationships. It gets compiled to regular JavaScript function calls.
""")
    add_paragraph(doc, """
Unidirectional Data Flow: React follows a unidirectional data flow, where data flows down from parent components to child components through props. This makes the application's data flow predictable and easier to debug.
""")
    add_paragraph(doc, """
Hooks: Introduced in React 16.8, hooks allow functional components to use state and other React features without writing class components. Hooks like useState, useEffect, and useContext provide powerful ways to manage component state and side effects.
""")
    
    # 7.2.4. React Ecosystem
    add_heading(doc, '7.2.4. React Ecosystem', 3)
    add_paragraph(doc, """
React Router: A popular routing library for React applications that allows developers to create single-page applications with multiple views. It provides declarative routing and supports nested routes, dynamic routing, and navigation guards.
""")
    add_paragraph(doc, """
Redux: A predictable state container for JavaScript applications, commonly used with React. Redux helps manage application state in a predictable way, making it easier to debug and test applications. It follows a unidirectional data flow and uses actions and reducers to update state.
""")
    add_paragraph(doc, """
React Query: A library for managing server state in React applications. It provides hooks for fetching, caching, synchronizing, and updating server data. React Query handles caching, background updates, and error handling automatically.
""")
    add_paragraph(doc, """
Material-UI and Ant Design: Popular UI component libraries that provide pre-built, customizable components for React applications. These libraries help developers build consistent and professional-looking user interfaces quickly.
""")
    
    # 7.2.5. Advantages of React
    add_heading(doc, '7.2.5. Advantages of React', 3)
    add_paragraph(doc, """
Performance: React's virtual DOM and efficient diffing algorithm make it highly performant. By minimizing direct DOM manipulation, React applications can handle complex UIs with smooth user interactions.
""")
    add_paragraph(doc, """
Reusability: React's component-based architecture promotes code reusability. Components can be easily shared across different parts of an application or even across different projects, reducing development time and improving maintainability.
""")
    add_paragraph(doc, """
Large Community: React has a large and active community of developers, which means extensive documentation, tutorials, and third-party libraries are available. This makes it easier to find solutions to common problems and learn best practices.
""")
    add_paragraph(doc, """
Developer Tools: React provides excellent developer tools, including React Developer Tools browser extension, which helps developers inspect component hierarchies, monitor state changes, and debug applications effectively.
""")
    add_paragraph(doc, """
Cross-Platform: React can be used to build web applications, mobile applications (React Native), and even desktop applications (Electron). This versatility makes it a valuable skill for developers working across different platforms.
""")
    
    # 7.3. Tailwind CSS Framework
    add_heading(doc, '7.3. Tailwind CSS Framework', 2)
    
    # 7.3.1. History of Tailwind CSS
    add_heading(doc, '7.3.1. History of Tailwind CSS', 3)
    add_paragraph(doc, """
Tailwind CSS was created by Adam Wathan and was first released in 2017. It was developed as a utility-first CSS framework to address the limitations of traditional CSS frameworks. Wathan, along with other developers, wanted to create a framework that would allow for rapid UI development without the constraints of pre-designed components.
""")
    add_paragraph(doc, """
The framework was initially developed as an internal tool for Wathan's consulting work. After seeing its effectiveness in speeding up development and creating consistent designs, he decided to open-source it. Tailwind CSS quickly gained popularity among developers due to its utility-first approach and the flexibility it provides.
""")
    add_paragraph(doc, """
Since its initial release, Tailwind CSS has grown significantly in popularity and has been adopted by many companies and developers worldwide. The framework has received regular updates and improvements, with new features and utilities being added to enhance developer experience and design capabilities.
""")
    
    # 7.3.2. Introduction to Tailwind CSS
    add_heading(doc, '7.3.2. Introduction to Tailwind CSS', 3)
    add_paragraph(doc, """
Tailwind CSS is a utility-first CSS framework that provides low-level utility classes to build custom designs without leaving your HTML. Unlike traditional CSS frameworks that provide pre-designed components, Tailwind CSS gives developers the building blocks to create any design they can imagine.
""")
    add_paragraph(doc, """
The framework is designed to be highly customizable and configurable. Developers can easily modify the default design system, add custom utilities, and extend the framework to match their project's specific needs. Tailwind CSS uses a mobile-first responsive design approach, making it easy to create responsive layouts.
""")
    add_paragraph(doc, """
Tailwind CSS is built with modern CSS features and best practices. It uses CSS Grid and Flexbox for layouts, CSS custom properties for theming, and modern CSS features for animations and transitions. The framework is also designed to be performant, with a focus on generating only the CSS that is actually used in the project.
""")
    
    # 7.3.3. Key Features of Tailwind CSS
    add_heading(doc, '7.3.3. Key Features of Tailwind CSS', 3)
    add_paragraph(doc, """
Utility-First Approach: Tailwind CSS provides utility classes for almost every CSS property, allowing developers to build designs directly in their HTML. This approach eliminates the need to write custom CSS for most common design patterns.
""")
    add_paragraph(doc, """
Responsive Design: Tailwind CSS includes responsive variants for all utility classes, making it easy to create responsive layouts. Developers can use prefixes like sm:, md:, lg:, and xl: to apply styles at different breakpoints.
""")
    add_paragraph(doc, """
Dark Mode Support: Tailwind CSS includes built-in support for dark mode, allowing developers to easily create dark and light themes for their applications. The framework provides dark: variants for applying styles in dark mode.
""")
    add_paragraph(doc, """
Customization: Tailwind CSS is highly customizable through its configuration file. Developers can modify colors, spacing, typography, breakpoints, and other design tokens to match their brand or design system.
""")
    add_paragraph(doc, """
PurgeCSS Integration: Tailwind CSS integrates with PurgeCSS to remove unused CSS in production, resulting in smaller file sizes and better performance. This ensures that only the CSS that is actually used is included in the final build.
""")
    
    # 7.3.4. Tailwind CSS Utilities
    add_heading(doc, '7.3.4. Tailwind CSS Utilities', 3)
    add_paragraph(doc, """
Layout Utilities: Tailwind CSS provides utilities for controlling layout, including display, position, top/right/bottom/left, and z-index. These utilities make it easy to create complex layouts without writing custom CSS.
""")
    add_paragraph(doc, """
Spacing Utilities: The framework includes utilities for margin, padding, width, height, and other spacing properties. These utilities use a consistent spacing scale that can be customized in the configuration.
""")
    add_paragraph(doc, """
Typography Utilities: Tailwind CSS provides utilities for controlling text color, font size, font weight, text alignment, and other typography properties. These utilities make it easy to create consistent typography across an application.
""")
    add_paragraph(doc, """
Background and Border Utilities: The framework includes utilities for background colors, background images, border styles, border colors, and border radius. These utilities provide comprehensive control over the visual appearance of elements.
""")
    add_paragraph(doc, """
Flexbox and Grid Utilities: Tailwind CSS provides utilities for Flexbox and CSS Grid layouts, making it easy to create complex layouts. These utilities include flex direction, justify content, align items, grid columns, and more.
""")
    
    # 7.3.5. Advantages of Tailwind CSS
    add_heading(doc, '7.3.5. Advantages of Tailwind CSS', 3)
    add_paragraph(doc, """
Rapid Development: Tailwind CSS allows developers to build user interfaces quickly by using utility classes directly in HTML. This eliminates the need to write custom CSS for most common design patterns, speeding up development time.
""")
    add_paragraph(doc, """
Consistency: The framework's design system ensures consistency across an application. By using predefined design tokens for colors, spacing, and typography, developers can create cohesive designs without much effort.
""")
    add_paragraph(doc, """
Maintainability: Tailwind CSS promotes maintainable code by reducing the amount of custom CSS needed. Since most styles are applied using utility classes, there's less custom CSS to maintain and debug.
""")
    add_paragraph(doc, """
Performance: The framework's PurgeCSS integration ensures that only the CSS that is actually used is included in the final build. This results in smaller file sizes and better performance for production applications.
""")
    add_paragraph(doc, """
Flexibility: Tailwind CSS provides the flexibility to create any design without being constrained by pre-designed components. Developers have full control over the appearance and behavior of their components.
""")
    
    # 8. DATASET DESCRIPTION
    add_heading(doc, '8. DATASET DESCRIPTION', 1)
    add_paragraph(doc, "The GrandCruise platform utilizes a relational database structure implemented through Django models. The system consists of four primary data models that form the foundation of the car booking platform.")
    update_section_8(doc)
    
    # 9. TESTING
    add_heading(doc, '9. TESTING', 1)
    update_section_9(doc)
    
    # 10. SOURCE CODE
    add_heading(doc, '10. SOURCE CODE', 1)
    update_section_10(doc)
    
    # 11. PROJECT SCREENSHOTS AND RESULTS
    add_heading(doc, '11. PROJECT SCREENSHOTS AND RESULTS', 1)
    add_paragraph(doc, """
This section presents visual documentation of the GrandCruise platform through screenshots and demonstrations of key features and functionalities. The screenshots showcase the user interface design, booking workflow, admin panel, and various system components that demonstrate the platform's capabilities and user experience.
""")
    
    add_heading(doc, '11.1 User Interface Screenshots', 2)
    add_paragraph(doc, """
The following screenshots demonstrate the modern and intuitive user interface of the GrandCruise platform, showcasing the responsive design and user-friendly navigation that enhances the overall booking experience.
""")
    add_paragraph(doc, '[Screenshots of Home Page, Car Listing, Car Details, and Booking Interface will be added here]')
    
    add_heading(doc, '11.2 Admin Panel Screenshots', 2)
    add_paragraph(doc, """
The admin panel provides comprehensive management capabilities for fleet administrators, including car management, booking oversight, user management, and system analytics.
""")
    add_paragraph(doc, '[Screenshots of Admin Dashboard, Car Management, Booking Management, and User Management will be added here]')
    
    add_heading(doc, '11.3 Booking Process Screenshots', 2)
    add_paragraph(doc, """
The booking workflow screenshots demonstrate the streamlined process from car selection to confirmation, highlighting the intuitive design and efficient user journey.
""")
    add_paragraph(doc, '[Screenshots of Booking Steps, Date Selection, Payment Interface, and Confirmation will be added here]')
    
    add_heading(doc, '11.4 Mobile Responsive Screenshots', 2)
    add_paragraph(doc, """
Mobile responsiveness is crucial for modern web applications. These screenshots demonstrate how the GrandCruise platform adapts to different screen sizes and maintains functionality across devices.
""")
    add_paragraph(doc, '[Screenshots of Mobile Interface, Tablet View, and Responsive Design will be added here]')
    
    # 12. CONCLUSION
    add_heading(doc, '12. CONCLUSION', 1)
    add_paragraph(doc, """
GrandCruise successfully demonstrates the implementation of a modern, full-stack car booking platform using Django and React, showcasing the power of combining robust backend technologies with dynamic frontend frameworks. The platform provides a seamless user experience for both customers and administrators, with comprehensive backend APIs built on Django REST Framework, responsive frontend interfaces developed with React and TypeScript, and advanced booking management capabilities that streamline the entire rental process.

The modular architecture ensures scalability and maintainability, allowing for future enhancements and integrations without compromising system performance. The integration of modern web technologies including Tailwind CSS for styling, JWT authentication for security, and comprehensive state management delivers a professional and user-friendly application that meets the demands of premium car rental services. The platform's database design with proper relationships, API endpoints with filtering and search capabilities, and responsive UI components create a cohesive ecosystem that enhances both user satisfaction and operational efficiency.

The implementation demonstrates best practices in full-stack development, including proper separation of concerns, RESTful API design, component-based architecture, and responsive design principles. The system successfully addresses the limitations of traditional car rental platforms by providing real-time availability updates, automated booking processes, and comprehensive admin tools for fleet management. GrandCruise represents a significant advancement in digital car rental solutions, setting new standards for user experience, system reliability, and business process automation in the premium mobility sector.
""")
    
    # 13. FUTURE WORK
    add_heading(doc, '13. FUTURE WORK', 1)
    add_paragraph(doc, """
Future enhancements for GrandCruise could include mobile application development using React Native or Flutter to provide native mobile experiences for both iOS and Android platforms, enabling users to book cars, manage reservations, and access real-time updates on-the-go. Payment gateway integration with services like Stripe, PayPal, or Square would enable secure online transactions, automated billing, and subscription-based rental models. Real-time notifications using WebSocket technology or push notifications would enhance user engagement by providing instant updates on booking confirmations, vehicle availability, and maintenance schedules.

Advanced analytics and reporting features could include business intelligence dashboards, revenue analytics, customer behavior insights, and predictive maintenance scheduling to optimize fleet utilization and improve operational efficiency. AI-powered recommendation systems could leverage machine learning algorithms to suggest personalized vehicle options based on user preferences, booking history, and seasonal demand patterns. Integration with third-party services such as GPS tracking systems, insurance providers, and maintenance service providers would create a comprehensive ecosystem for enhanced user experience and streamlined business operations.

Additional future developments could encompass blockchain-based smart contracts for automated rental agreements, IoT integration for vehicle monitoring and diagnostics, virtual reality showrooms for immersive car viewing experiences, and integration with ride-sharing platforms for hybrid mobility solutions. These enhancements would position GrandCruise as a cutting-edge platform in the evolving landscape of digital mobility services.
""")
    
    # 14. REFERENCES
    add_heading(doc, '14. REFERENCES', 1)
    add_paragraph(doc, '1. Django Documentation. https://docs.djangoproject.com/')
    add_paragraph(doc, '2. React Documentation. https://react.dev/')
    add_paragraph(doc, '3. Tailwind CSS Documentation. https://tailwindcss.com/docs')
    add_paragraph(doc, '4. Django REST Framework Documentation. https://www.django-rest-framework.org/')
    add_paragraph(doc, '5. MySQL Documentation. https://dev.mysql.com/doc/')
    add_paragraph(doc, '6. Web Development Best Practices. Modern Web Development with React and Django.')
    add_paragraph(doc, '[1] Waspodo, Bayu, Qurrotul Aini, and Syamsuri Nur. "Development of car rental management information system." In Proceeding International Conference on Information Systems For Business Competitiveness (ICISBC), pp. 101-105. 2011.')
    add_paragraph(doc, '[2] Osman, Mohd Nizam, Nurzaid Md Zain, Zulfikri Paidi, Khairul Anwar Sedek, Mohamad NajmuddinYusoff, and Mushahadah Maghribi. "Online Car Rental System Using Web-Based and SMS Technology." Computing Research & Innovation (CRINN) 2 (2017): 277.')
    add_paragraph(doc, '[3] Fink, Andreas, and Torsten Reiners. "Modeling and solving the short-term car rental logistics problem." Transportation Research Part E: Logistics and Transportation Review 42, no. 4 (2006): 272-292.')
    add_paragraph(doc, '[4] Khaled, Mr Shah Mostafa, Shamsil Arefin, Datta Sree Rajib Kumar, and Ariful Hossain Tuhin. "Software Requirements Specification for Online Car Rental System." (2015).')
    add_paragraph(doc, '[5] Carroll, William J., and Richard C. Grimes. "Evolutionary change in product management: Experiences in the car rental industry." Interfaces 25, no. 5 (1995): 84-104.')
    add_paragraph(doc, '[6] Beck, Kent, Mike Beedle, Arie Van Bennekum, Alistair Cockburn, Ward Cunningham, Martin Fowler, James Grenning et al. "Manifesto for agile software development." (2001): 2006.')
    add_paragraph(doc, '[7] Abrahamsson, Pekka, Outi Salo, Jussi Ronkainen, and Juhani Warsta. "Agile software development methods: Review and analysis." arXiv preprint arXiv:1709.08439 (2017).')
    add_paragraph(doc, '[8] Thakur, A., & Dhiman, K. (2021). Chat Room Using HTML, PHP, CSS, JS, AJAX. International Research Journal of Engineering and Technology (IRJET), 08(June), 1948–1951. https://doi.org/https://doi.org/10.6084/m9.figshare.14869167')
    add_paragraph(doc, '[9] Thakur, Amey and Karan Dhiman. "Chat Room Using HTML, PHP, CSS, JS, AJAX." ArXiv abs/2106.14704 (2021): n. pag.')
    add_paragraph(doc, '[10] Soares, Hécio A., and Raimundo S. Moura. "A methodology to guide writing Software Requirements Specification document." In 2015 Latin American Computing Conference (CLEI), pp. 1-11. IEEE, 2015.')

    doc.save('Project_Documentation.docx')

if __name__ == '__main__':
    main() 